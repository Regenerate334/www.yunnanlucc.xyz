# -*- coding: utf-8 -*-
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_markdown_tasks(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    tasks = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|') or '---' in line or '序号' in line:
            continue
        parts = [p.strip() for p in line.split('|')[1:-1]]
        if len(parts) >= 8:
            tasks.append({
                'index': parts[0],
                'id': parts[1],
                'category': parts[2],
                'difficulty': parts[3],
                'question': parts[4],
                'tools': parts[5],
                'args': parts[6],
                'baseline': parts[7]
            })
    return tasks

def create_tasks_docx(tasks, output_path):
    doc = Document()
    
    # Page setup - Landscape for better table fit
    section = doc.sections[0]
    section.page_width = Inches(11.69) # A4 Landscape
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("附表：GeoAI Agent 自然语言评价任务集（共 72 题）")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtitle / Description
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_run = sub.add_run("说明：本任务集基于国土空间规划与土地利用监测全业务链构建，涵盖 6 类业务场景与 3 个认知难度层级，采用 6×3×4 = 72 题正交分层设计。每项任务均明确标定期望工具调用、关键参数与标准数据库抽取口径/政策事实依据，用于客观定量评估 Agent 的工具调度、参数解析与结果一致性。")
    sub_run.font.name = "宋体"
    sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Create Table
    headers = ["序号", "任务类别", "难度", "自然语言测试问题（Prompt）", "期望工具", "关键参数", "基准结果/政策参考抽取口径"]
    col_widths = [Inches(0.5), Inches(1.3), Inches(0.8), Inches(3.2), Inches(1.3), Inches(1.5), Inches(1.5)]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1F497D") # Navy Blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    # Set header row repeat on every page
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    # Add Data Rows
    for row_idx, task in enumerate(tasks):
        row_cells = table.add_row().cells
        data = [
            task['index'],
            task['category'],
            task['difficulty'],
            task['question'],
            task['tools'],
            task['args'],
            task['baseline']
        ]
        
        bg_color = "F2F5F9" if row_idx % 2 == 1 else "FFFFFF"
        
        for col_idx, text in enumerate(data):
            cell = row_cells[col_idx]
            cell.text = text
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            p = cell.paragraphs[0]
            if col_idx in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif col_idx in [1, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            for r in p.runs:
                r.font.name = "宋体"
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(40, 40, 40)
    
    # Apply simple borders to table
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="1F497D"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="1F497D"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)
    
    doc.save(output_path)
    print(f"Successfully generated docx: {output_path}")

if __name__ == '__main__':
    md_file = r"c:\projects\webgis\my_webgis_project\ops\ai\evaluation\tasks\task_design_72.md"
    out_dir = r"c:\projects\webgis\my_webgis_project\ops\ai\evaluation\reports"
    os.makedirs(out_dir, exist_ok=True)
    out_file1 = os.path.join(out_dir, "GeoAI_Agent_72_Evaluation_Tasks.docx")
    out_file2 = r"C:\Users\DIY\GeoAI_Agent_72_Evaluation_Tasks.docx"
    
    tasks = parse_markdown_tasks(md_file)
    print(f"Parsed {len(tasks)} tasks from {md_file}")
    
    create_tasks_docx(tasks, out_file1)
    create_tasks_docx(tasks, out_file2)
